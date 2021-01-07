import docx
from docx.shared import Cm
from docx.oxml.table import CT_Row, CT_Tc
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


import datetime
import locale
import os.path
import sys
from copy import deepcopy
from io import BytesIO

import pytils
from reportlab.lib import colors
from reportlab.lib.colors import black
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, A5, landscape, portrait
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import NextPageTemplate, Indenter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Frame, PageTemplate, FrameBreak, Table, TableStyle, PageBreak

from api.patients.views import research_last_result_every_month
from appconf.manager import SettingManager
from clients.models import Card, DispensaryReg, DispensaryRegPlans
from directory.models import DispensaryPlan, Researches
from hospitals.models import Hospitals
from laboratory.settings import FONTS_FOLDER
from laboratory.utils import strdate
from statistics_tickets.models import VisitPurpose
from utils.dates import normalize_date
from utils.flowable import InteractiveListBoxField, InteractiveTextField, InteractiveListTypeMedExam
from laboratory.utils import strfdatetime


# def form_01(request_data):
#     """
#     Форма 025/у - титульный лист амбулаторной карты
#     Приказ Минздрава России от 15.12.2014 N 834н (ред. от 09.01.2018)
#     http://docs.cntd.ru/document/436733768) Об утверждении критериев оценки качества медицинской помощи
#     ПРИКАЗ Минздрава России от 10 мая 2017 года N 203н https://minjust.consultant.ru/documents/35361
#     """
#     ind_card = Card.objects.get(pk=request_data["card_pk"])
#     patient_data = ind_card.get_data_individual()
#     hospital: Hospitals = request_data["hospital"]
#
#     hospital_name = hospital.safe_short_title
#     hospital_address = hospital.safe_address
#     hospital_kod_ogrn = hospital.safe_ogrn
#
#     if sys.platform == 'win32':
#         locale.setlocale(locale.LC_ALL, 'rus_rus')
#     else:
#         locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')
#
#     pdfmetrics.registerFont(TTFont('PTAstraSerifBold', os.path.join(FONTS_FOLDER, 'PTAstraSerif-Bold.ttf')))
#     pdfmetrics.registerFont(TTFont('PTAstraSerifReg', os.path.join(FONTS_FOLDER, 'PTAstraSerif-Regular.ttf')))
#
#     buffer = BytesIO()
#     doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=25 * mm, rightMargin=5 * mm, topMargin=6 * mm, bottomMargin=6 * mm, allowSplitting=1, title="Форма {}".format("025/у"))
#     width, height = portrait(A4)
#     styleSheet = getSampleStyleSheet()
#     style = styleSheet["Normal"]
#     style.fontName = "PTAstraSerifReg"
#     style.fontSize = 10
#     style.leading = 12
#     style.spaceAfter = 0.5 * mm
#     styleBold = deepcopy(style)
#     styleBold.fontName = "PTAstraSerifBold"
#     styleCenter = deepcopy(style)
#     styleCenter.alignment = TA_CENTER
#     styleCenter.fontSize = 12
#     styleCenter.leading = 15
#     styleCenter.spaceAfter = 1 * mm
#     styleCenterBold = deepcopy(styleBold)
#     styleCenterBold.alignment = TA_CENTER
#     styleCenterBold.fontSize = 12
#     styleCenterBold.leading = 15
#     styleCenterBold.face = 'PTAstraSerifBold'
#     styleCenterBold.borderColor = black
#     styleJustified = deepcopy(style)
#     styleJustified.alignment = TA_JUSTIFY
#     styleJustified.spaceAfter = 4.5 * mm
#     styleJustified.fontSize = 12
#     styleJustified.leading = 4.5 * mm
#
#     objs = []
#
#     styleT = deepcopy(style)
#     styleT.alignment = TA_LEFT
#     styleT.fontSize = 10
#     styleT.leading = 4.5 * mm
#     styleT.face = 'PTAstraSerifReg'
#
#     print_district = ''
#     if SettingManager.get("district", default='True', default_type='b'):
#         if ind_card.district is not None:
#             print_district = 'Уч: {}'.format(ind_card.district.title)
#
#     opinion = [
#         [
#             Paragraph('<font size=11>{}<br/>Адрес: {}<br/>ОГРН: {} <br/><u>{}</u> </font>'.format(hospital_name, hospital_address, hospital_kod_ogrn, print_district), styleT),
#             Paragraph('<font size=9 >Код формы по ОКУД:<br/>Код организации по ОКПО: 31348613<br/>' 'Медицинская документация<br/>Учетная форма № 025/у</font>', styleT),
#         ],
#     ]
#
#     tbl = Table(opinion, 2 * [90 * mm])
#     tbl.setStyle(
#         TableStyle(
#             [
#                 ('GRID', (0, 0), (-1, -1), 0.75, colors.white),
#                 ('LEFTPADDING', (1, 0), (-1, -1), 80),
#                 ('VALIGN', (0, 0), (-1, -1), 'TOP'),
#             ]
#         )
#     )
#
#     objs.append(tbl)
#     space_symbol = '&nbsp;'
#     if patient_data['age'] < SettingManager.get("child_age_before", default='15', default_type='i'):
#         patient_data['serial'] = patient_data['bc_serial']
#         patient_data['num'] = patient_data['bc_num']
#     else:
#         patient_data['serial'] = patient_data['passport_serial']
#         patient_data['num'] = patient_data['passport_num']
#
#     p_phone = ''
#     if patient_data['phone']:
#         p_phone = 'тел. ' + ", ".join(patient_data['phone'])
#
#     card_num_obj = patient_data['card_num'].split(' ')
#     p_card_num = card_num_obj[0]
#     if len(card_num_obj) == 2:
#         p_card_type = '(' + str(card_num_obj[1]) + ')'
#     else:
#         p_card_type = ''
#     content_title = [
#         Indenter(left=0 * mm),
#         Spacer(1, 1 * mm),
#         Paragraph('МЕДИЦИНСКАЯ КАРТА ПАЦИЕНТА, <br/> ПОЛУЧАЮЩЕГО МЕДИЦИНСКУЮ ПОМОЩЬ В АМБУЛАТОРНЫХ УСЛОВИЯХ', styleCenter),
#         Paragraph(
#             '{}<font size=14>№</font><font fontname="PTAstraSerifBold" size=17> <u>{}</u></font><font size=14> {}</font>'.format(3 * space_symbol, p_card_num, p_card_type), styleCenter
#         ),
#         Spacer(1, 2 * mm),
#         Paragraph('1.Дата заполнения медицинской карты: {}'.format(pytils.dt.ru_strftime(u"%d %B %Y", inflected=True, date=datetime.datetime.now())), style),
#         Paragraph("2. Фамилия, имя, отчество:&nbsp;  <font size=11.7 fontname ='PTAstraSerifBold'> {} </font> ".format(patient_data['fio']), style),
#         Paragraph('3. Пол: {} {} 4. Дата рождения: {}'.format(patient_data['sex'], 3 * space_symbol, patient_data['born']), style),
#         Paragraph('5. Место регистрации: {}'.format(patient_data['main_address']), style),
#         Paragraph('{}'.format(p_phone), style),
#         Paragraph('6. Местность: городская — 1, сельская — 2', style),
#         Paragraph(
#             '7. Полис ОМС: серия {} №: {} {}' '8. СНИЛС: {}'.format(patient_data['oms']['polis_serial'], patient_data['oms']['polis_num'], 13 * space_symbol, patient_data['snils']), style
#         ),
#         Paragraph('9. Наименование страховой медицинской организации: {}'.format(patient_data['oms']['polis_issued']), style),
#         Paragraph(
#             '10. Код категории льготы: {} {} 11. Документ: {} &nbsp; серия: {} &nbsp;&nbsp; №: {}'.format(
#                 8 * space_symbol, 25 * space_symbol, patient_data['type_doc'], patient_data['serial'], patient_data['num']
#             ),
#             style,
#         ),
#     ]
#
#     objs.extend(content_title)
#     if SettingManager.get("dispensary", default='True', default_type='b'):
#         objs.append(Paragraph('12. Заболевания, по поводу которых осуществляется диспансерное наблюдение:', style))
#         objs.append(Spacer(1, 2 * mm))
#
#         styleTCenter = deepcopy(styleT)
#         styleTCenter.alignment = TA_CENTER
#         styleTCenter.leading = 3.5 * mm
#
#         opinion = [
#             [
#                 Paragraph('<font size=9>Дата начала диспансерного наблюдения </font>', styleTCenter),
#                 Paragraph('<font size=9 >Дата прекращения диспансерного наблюдения</font>', styleTCenter),
#                 Paragraph('<font size=9 >Диагноз</font>', styleTCenter),
#                 Paragraph('<font size=9 >Код по МКБ-10</font>', styleTCenter),
#                 Paragraph('<font size=9 >Врач</font>', styleTCenter),
#             ],
#         ]
#         for i in range(0, 5):
#             para = ['', '', '', '', '']
#             opinion.append(para)
#
#         row_height = []
#         for i in opinion:
#             row_height.append(6 * mm)
#
#         row_height[0] = None
#
#         tbl = Table(opinion, colWidths=(27 * mm, 30 * mm, 75 * mm, 20 * mm, 27 * mm), rowHeights=row_height)
#
#         tbl.setStyle(
#             TableStyle(
#                 [
#                     ('GRID', (0, 0), (-1, -1), 1.0, colors.black),
#                     ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
#                 ]
#             )
#         )
#
#         objs.append(tbl)
#
#     def first_pages(canvas, document):
#         canvas.saveState()
#         canvas.restoreState()
#
#     def later_pages(canvas, document):
#         canvas.saveState()
#         canvas.restoreState()
#
#     doc.build(objs, onFirstPage=first_pages, onLaterPages=later_pages)
#
#     pdf = buffer.getvalue()
#     buffer.seek(0)
#
#     today = datetime.datetime.now()
#     date_now1 = datetime.datetime.strftime(today, "%y%m%d%H%M%S%f")[:-3]
#     # date_now_str = str(client_id) + str(date_now1)
#     # date_now_str = str(client_id) + str(date_now1)
#     date_now_str = str(date_now1)
#     dir_param = SettingManager.get("dir_param", default='/tmp', default_type='s')
#     file_dir = os.path.join(dir_param, date_now_str + '_dir.pdf')
#
#     def save(form, filename: str):
#         with open(filename, 'wb') as f:
#             f.write(form.read())
#
#     save(buffer, filename=file_dir)
#
#     from pdf2docx import Converter
#
#     # pdf_file = '/Users/sergejkasanenko/Documents/tmp/results.pdf'
#     docx_file = os.path.join(dir_param, date_now_str + '_dir.docx')
#
#     # convert pdf to docx
#     cv = Converter(file_dir)
#     cv.convert(docx_file, start=0, end=None)
#     cv.close()
#     buffer.close()
#     os.remove(file_dir)
#
#     doc = docx.Document(docx_file)
#     os.remove(docx_file)
#
#     buffer.close()
#
#     return doc
#
#
# def form_03(request_data):
#     doc = docx.Document()
#     sections= doc.sections[-1].start_type
#     sections.left_margin = Cm(1.0)
#     sections.right_margin = Cm(1.0)
#
#     doc.add_paragraph('Министерство здравоохранения Российской Федерации\nОГАУЗ «Иркутская медико-санитарная часть №2»'
#                       '\nг. Иркутск, ул. Байкальская,  201', style='No Spacing')
#     doc.add_paragraph('', style='No Spacing')
#     doc.add_paragraph('', style='No Spacing')
#
#     table = doc.add_table(rows=1, cols=14)
#     table.style = 'Table Grid'
#     kod_ogrn = "1033801542576"
#     for col in range(14):
#         cell = table.cell(0, col)
#         cell.width = Cm(0.5)
#         if col == 0:
#             cell.width = Cm(2.3)
#             cell.text = "Код ОГРН"
#         else:
#             cell.text = kod_ogrn[col-1]
#
#     tbl = table._tbl
#     # http://officeopenxml.com/WPtableBorders.php
#
#     def change_border(tbl):
#         for cell in tbl.iter_tcs():
#             tcPr = cell.tcPr
#             tcBorders = OxmlElement("w:tcBorders")
#             top = OxmlElement("w:top")
#             top.set(qn("w:val"), "nil")
#
#             left = OxmlElement("w:left")
#             left.set(qn("w:val"), "nil")
#
#             bottom = OxmlElement("w:bottom")
#             bottom.set(qn("w:val"), "nil")
#
#             right = OxmlElement("w:right")
#             right.set(qn("w:val"), "nil")
#
#             tcBorders.append(top)
#             tcBorders.append(left)
#             tcBorders.append(bottom)
#             tcBorders.append(right)
#             tcPr.append(tcBorders)
#             break
#
#     change_border(tbl)
#
#     set_pt = Pt(12)
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.add_run('ЗАКЛЮЧИТЕЛЬНЫЙ АКТ \n по  результатам  проведенного  периодического  медицинского  осмотра \n'
#                       '(обследования) работников за 2020 год.').bold = True
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.add_run('12.10.2020').bold = True
#
#     paragraph = doc.add_paragraph()
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('Комиссией в составе:')
#
#     paragraph = doc.add_paragraph('')
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('Председатель  врачебной  комиссии:')
#     paragraph = doc.add_paragraph('')
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('Врач-терапевт       :')
#
#     paragraph = doc.add_paragraph('')
#     paragraph = doc.add_paragraph('')
#     paragraph.style = 'No Spacing'
#     paragraph.add_run(
#         'Составлен настоящий акт по результатам  проведенного  периодического  медицинского  осмотра (обследования) '
#         'работников  МОУ  ИРМО  « Кыцигировская НШДС »   в период с 16.06.2020 по 25.06.2020.')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.add_run('1. Число работников организации (предприятия), цеха:')
#     paragraph.style = 'No Spacing'
#     table = doc.add_table(rows=2, cols=2)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(14)
#     table.cell(0, 1).width = Cm(3)
#     table.cell(1, 0).width = Cm(14)
#     table.cell(1, 1).width = Cm(3)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.add_run('2. Число работников организации (предприятия), цеха, работающих с вредными и (или) опасными '
#                       'веществами и производственными факторами, а так же на работах*:')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     paragraph.style = 'No Spacing'
#     table = doc.add_table(rows=2, cols=2)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(14)
#     table.cell(0, 1).width = Cm(3)
#     table.cell(1, 0).width = Cm(14)
#     table.cell(1, 1).width = Cm(3)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('3. Число работников, подлежащих периодическому медицинскому осмотру (обследованию), '
#                       'работающих в контакте с вредными и (или) опасными веществами и производственными факторами, '
#                       'а так же на работах* в данном году:')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=2, cols=2)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(14)
#     table.cell(0, 1).width = Cm(3)
#     table.cell(1, 0).width = Cm(14)
#     table.cell(1, 1).width = Cm(3)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('4. Число работников, прошедших периодический медицинский осмотр (обследования):')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=2, cols=2)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(14)
#     table.cell(0, 1).width = Cm(3)
#     table.cell(1, 0).width = Cm(14)
#     table.cell(1, 1).width = Cm(3)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.add_run('Поименный список работников, прошедших периодический медицинский осмотр (обследования): ')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=1, cols=3)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(1)
#     table.cell(0, 1).width = Cm(10)
#     table.cell(0, 2).width = Cm(6)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('5. % охвата периодическими медицинскими осмотрами:')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=2, cols=2)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(14)
#     table.cell(0, 1).width = Cm(3)
#     table.cell(1, 0).width = Cm(14)
#     table.cell(1, 1).width = Cm(3)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('6. Число работников, не завершивших периодический медицинский осмотр (обследования):')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=2, cols=2)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(14)
#     table.cell(0, 1).width = Cm(3)
#     table.cell(1, 0).width = Cm(14)
#     table.cell(1, 1).width = Cm(3)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('Поименный список работников, не завершивших периодический медицинский осмотр (обследования):')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=1, cols=3)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(1)
#     table.cell(0, 1).width = Cm(10)
#     table.cell(0, 2).width = Cm(6)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('7. Число работников, не прошедших периодический медицинский осмотр (обследование):')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=2, cols=2)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(14)
#     table.cell(0, 1).width = Cm(3)
#     table.cell(1, 0).width = Cm(14)
#     table.cell(1, 1).width = Cm(3)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('Поименный список работников, не прошедших периодический медицинский осмотр (обследование):')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#     table = doc.add_table(rows=1, cols=4)
#     table.style = 'Table Grid'
#     table.cell(0, 0).width = Cm(1)
#     table.cell(0, 1).width = Cm(8)
#     table.cell(0, 2).width = Cm(6)
#     table.cell(0, 3).width = Cm(2)
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('8. Заключение по результатам данного периодического медицинского осмотра (обследования)')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#
#     paragraph = doc.add_paragraph()
#     paragraph.paragraph_format.space_before = set_pt
#     paragraph.style = 'No Spacing'
#     paragraph.add_run('8.1 Сводная таблица N 1:')
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#
#     table = doc.add_table(rows=13, cols=3)
#     table.style = 'Table Grid'
#     for i in range(13):
#         table.cell(i, 0).width = Cm(13)
#         table.cell(i, 1).width = Cm(2)
#         table.cell(i, 1).text = "Всего"
#         table.cell(i, 2).width = Cm(3)
#         table.cell(i, 2).text = "В том числе женщин"
#
#     section = doc.sections[0]
#     section.left_margin = Cm(3)
#     section.top_margin = Cm(1.5)
#     section.right_margin = Cm(1.5)
#
#     return doc


def form_04(request_data):
    hospital: Hospitals = request_data["hospital"]

    hospital_name = hospital.safe_short_title
    hospital_address = hospital.safe_address
    hospital_kod_ogrn = hospital.safe_ogrn

    if sys.platform == 'win32':
        locale.setlocale(locale.LC_ALL, 'rus_rus')
    else:
        locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')

    pdfmetrics.registerFont(TTFont('PTAstraSerifBold', os.path.join(FONTS_FOLDER, 'PTAstraSerif-Bold.ttf')))
    pdfmetrics.registerFont(TTFont('PTAstraSerifReg', os.path.join(FONTS_FOLDER, 'PTAstraSerif-Regular.ttf')))

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=25 * mm, rightMargin=5 * mm, topMargin=6 * mm, bottomMargin=6 * mm, allowSplitting=1, title="Форма {}".format("025/у"))
    width, height = portrait(A4)
    styleSheet = getSampleStyleSheet()
    style = styleSheet["Normal"]
    style.fontName = "PTAstraSerifReg"
    style.fontSize = 12
    style.leading = 15
    style.spaceAfter = 0.5 * mm
    styleBold = deepcopy(style)
    styleBold.fontName = "PTAstraSerifBold"
    styleCenter = deepcopy(style)
    styleCenter.alignment = TA_CENTER
    styleCenter.fontSize = 12
    styleCenter.leading = 7
    styleCenter.spaceAfter = 1 * mm
    styleCenterBold = deepcopy(styleBold)
    styleCenterBold.alignment = TA_CENTER
    styleCenterBold.fontSize = 12
    styleCenterBold.leading = 15
    styleCenterBold.face = 'PTAstraSerifBold'
    styleCenterBold.borderColor = black
    styleJustified = deepcopy(style)
    styleJustified.alignment = TA_JUSTIFY
    styleJustified.spaceAfter = 4.5 * mm
    styleJustified.fontSize = 12
    styleJustified.leading = 4.5 * mm

    objs = []

    styleT = deepcopy(style)
    styleT.alignment = TA_LEFT
    styleT.fontSize = 11
    styleT.leading = 4.5 * mm
    styleT.face = 'PTAstraSerifReg'

    styleTCenter = deepcopy(styleT)
    styleTCenter.alignment = TA_CENTER

    opinion = [
        [
            Paragraph('<font size=11>{}<br/>Адрес: {}<br/></font>'.format(hospital_name, hospital_address), styleT),
            Paragraph('', styleT),
        ],
    ]

    tbl = Table(opinion, 2 * [90 * mm])
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.white),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )

    objs.append(tbl)
    opinion = [
        [
            Paragraph('Код ОГРН', styleT),
            Paragraph(f"{hospital_kod_ogrn[0]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[1]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[2]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[3]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[4]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[5]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[6]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[7]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[8]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[9]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[10]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[11]}", styleT),
            Paragraph(f"{hospital_kod_ogrn[12]}", styleT),
        ],
    ]
    col_width = [6 * mm for i in range(13)]
    col_width.insert(0, 22 * mm)
    tbl = Table(opinion, hAlign='LEFT', rowHeights=6 * mm, colWidths=tuple(col_width))
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (1, 0), (-1, 0), 0.75, colors.black),
                ('LINEABOVE', (0, 0), (0, -1), 0.5, colors.white),
                ('LINEBEFORE', (0, 0), (0, -1), 2.5, colors.white),
                ('LINEBELOW', (0, 0), (0, -1), 1.5, colors.white),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3.5 * mm),
                ('LEFTPADDING', (1, 0), (-1, -1), 1 * mm),
            ]
        )
    )

    objs.append(Spacer(1, 3 * mm))
    objs.append(tbl)

    space_symbol = '&nbsp;'
    objs.append(Spacer(1, 13 * mm))
    objs.append(Paragraph('ЗАКЛЮЧИТЕЛЬНЫЙ АКТ', styleCenterBold))
    objs.append(Paragraph('по  результатам  проведенного  периодического  медицинского  осмотра ', styleCenterBold))
    objs.append(Paragraph('(обследования) работников за 2020 год. ', styleCenterBold))
    objs.append(Spacer(1, 4 * mm))

    objs.append(Paragraph('12.10. 2020 г.', styleBold))
    objs.append(Spacer(1, 2 * mm))
    objs.append(Paragraph('Комиссией в составе', style))
    objs.append(Paragraph(f'Председатель  врачебной  комиссии {space_symbol * 10} {hospital_name} Кирилюк К.В.', style))
    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph(f'Врач-терапевт {space_symbol * 10} {hospital_name} Шатурская Л.Е.', style))
    objs.append(
        Paragraph(
            f'Составлен настоящий акт по результатам  проведенного  периодического  медицинского  осмотра (обследования) работников  МОУ  ИРМО  « Кыцигировская НШДС »  '
            f' в период с 16.06.2020 по 25.06.2020.',
            style,
        )
    )

    # col_width = (150 * mm, 20 * mm)
    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('1. Число работников организации (предприятия), цеха:', style))
    objs = all_and_women(objs, styleT)

    objs.append(Spacer(1, 3 * mm))
    objs.append(
        Paragraph('2. Число работников организации (предприятия), цеха, работающих с вредными и (или) опасными веществами и производственными факторами, а так же на работах*:', style)
    )
    objs = all_and_women(objs, styleT)

    objs.append(Spacer(1, 3 * mm))
    objs.append(
        Paragraph(
            '3. Число работников, подлежащих периодическому медицинскому осмотру (обследованию), работающих в контакте с вредными и (или) опасными веществами и '
            'производственными факторами, а так же на работах* в данном году:',
            style,
        )
    )
    objs = all_and_women(objs, styleT)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('4. Число работников, прошедших периодический медицинский осмотр (обследования):', style))
    objs = all_and_women(objs, styleT)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('Поименный список работников, завершивших периодический медицинский осмотр ', style))
    opinion = [
        [
            Paragraph('№', styleT),
            Paragraph('Ф.И.О.', styleT),
            Paragraph('пол', styleT),
            Paragraph('Дата рождения', styleT),
            Paragraph('Должность', styleT),
            Paragraph('Заключение', styleT),
        ],
        [
            Paragraph('', styleT),
            Paragraph('', styleT),
            Paragraph('', styleT),
            Paragraph('', styleT),
            Paragraph('', styleT),
            Paragraph('', styleT),
        ],
    ]
    tbl = Table(
        opinion,
        colWidths=(7 * mm, 80 * mm, 10 * mm, 25 * mm, 30 * mm, 25 * mm, ),
    )
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('5. % охвата периодическими медицинскими осмотрами:', style))
    objs = all_and_women(objs, styleT)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('6. Число работников, не завершивших периодический медицинский осмотр (обследования):', style))
    objs = all_and_women(objs, styleT)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('Поименный список работников, не завершивших периодический медицинский осмотр (обследования):', style))
    opinion = [
        [
            Paragraph('№', styleT),
            Paragraph('Ф.И.О.', styleT),
            Paragraph('Подразделение предприятия', styleT),
        ],
        [
            Paragraph('', styleT),
            Paragraph('', styleT),
            Paragraph('', styleT),
        ],
    ]
    tbl = Table(
        opinion,
        colWidths=(
            7 * mm,
            120 * mm,
            50 * mm,
        ),
    )
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('7. Число работников, не прошедших периодический медицинский осмотр (обследование):', style))
    opinion = [
        [
            Paragraph('всего,', styleT),
            Paragraph('18', styleT),
        ],
        [
            Paragraph('в том числе женщин', styleT),
            Paragraph('-', styleT),
        ],
        [
            Paragraph('в том числе по причине: (медосмотр прошел при переводе, после лечения)', styleT),
            Paragraph('-', styleT),
        ],
        [
            Paragraph('больничный лист', styleT),
            Paragraph('-', styleT),
        ],
        [
            Paragraph('командировка', styleT),
            Paragraph('-', styleT),
        ],
        [
            Paragraph('очередной отпуск', styleT),
            Paragraph('-', styleT),
        ],
        [
            Paragraph('увольнение', styleT),
            Paragraph('-', styleT),
        ],
        [
            Paragraph('отказ от прохождения', styleT),
            Paragraph('-', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(150 * mm, 20 * mm))
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('Поименный список работников, не прошедших периодический медицинский осмотр (обследование):', style))
    opinion = [
        [
            Paragraph('№', styleT),
            Paragraph('Ф.И.О.', styleT),
            Paragraph('Подразделение предприятия', styleT),
            Paragraph('Причина', styleT),
        ],
        [
            Paragraph('', styleT),
            Paragraph('', styleT),
            Paragraph('', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(7 * mm, 80 * mm, 60 * mm, 30 * mm,),)
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 3 * mm))
    objs.append(Paragraph('8. Заключение по результатам данного периодического медицинского осмотра (обследования)', style))
    objs.append(Spacer(1, 1 * mm))
    objs.append(Paragraph('8.1 Сводная таблица N 1:', style))
    opinion = [
        [
            Paragraph('Результаты периодического медицинского осмотра (обследования)', styleT),
            Paragraph('Всего', styleT),
            Paragraph('В том числе женщин', styleT),
        ],
        [
            Paragraph('Число лиц, профпригодных к работе с вредными и (или) опасными веществами и производственными факторами, к видам работ*', styleT),
            Paragraph('-', styleT),
            Paragraph('-', styleT),
        ],
        [
            Paragraph('Число лиц, временно профнепригодных к работе с вредными и (или) опасными веществами и производственными факторами, к видам работ*', styleT),
            Paragraph('-', styleT),
            Paragraph('-', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(120 * mm, 27 * mm, 27 * mm, ),)
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs = add_needs_text(objs, 'Число лиц, постоянно профнепригодных к работе с вредными и (или) опасными веществами и производственными факторами, к видам работ*', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц нуждающихся в дообследовании (заключение не дано)', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц с подозрением на профессиональное заболевание', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц, нуждающихся в обследовании в центре профпатологии', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Люди', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц, нуждающихся в амбулаторном обследовании и лечении', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц, нуждающихся в стационарном обследовании и лечении: (оперативное лечение)', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц, нуждающихся в санаторно-курортном лечении', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц, нуждающихся в лечебно-профилактическом питании', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц, нуждающихся в диспансерном наблюдении', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs = add_needs_text(objs, 'Число лиц, нуждающихся в направлении на медико-социальную экспертизу', styleT)
    objs = add_fio_spec_diagnoz(objs, styleT)

    objs.append(Spacer(1, 5 * mm))
    objs.append(Paragraph('8.3 Выявлено лиц с подозрением на профессиональное заболевание:', style))
    opinion = [
        [
            Paragraph('Nп/п', styleT),
            Paragraph('Ф.И.О.', styleT),
            Paragraph('Подразделение предприятия', styleT),
            Paragraph('Профессия, должность', styleT),
            Paragraph('Вредные и (или) опасные вещества и производственные факторы', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(10 * mm, 90 * mm, 25 * mm, 25 * mm, 25 * mm,), )
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 5 * mm))
    objs.append(Paragraph('8.4 Выявлено впервые в жизни хронических соматических заболеваний:', style))
    opinion = [
        [
            Paragraph('№', styleT),
            Paragraph('Класс заболевания по МКБ-10', styleT),
            Paragraph('Количество работников (всего)', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(10 * mm, 130 * mm, 30 * mm,), )
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 5 * mm))
    objs.append(Paragraph('8.5 Выявлено впервые в жизни хронических профессиональных заболеваний:', style))
    opinion = [
        [
            Paragraph('№', styleT),
            Paragraph('Класс заболевания по МКБ-10', styleT),
            Paragraph('Количество работников (всего)', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(10 * mm, 130 * mm, 30 * mm,), )
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 5 * mm))
    objs.append(Paragraph('9. Результаты выполнения рекомендаций предыдущего заключительного акта. по результатам проведенного периодического медицинского осмотра (обследования) работников.', style))

    opinion = [
        [
            Paragraph('Nп/п', styleT),
            Paragraph('Мероприятия', styleT),
            Paragraph('Подлежало (чел.)', styleT),
            Paragraph('абс.', styleT),
            Paragraph('в %', styleT),
        ],
        [
            Paragraph('1', styleT),
            Paragraph('Обследование в центре профпатологии', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
        [
            Paragraph('2', styleT),
            Paragraph('Дообследование', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
        [
            Paragraph('3', styleT),
            Paragraph('Лечение и обследование амбулаторное', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
        [
            Paragraph('4', styleT),
            Paragraph('Лечение и обследование стационарное', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
        [
            Paragraph('5', styleT),
            Paragraph('Санаторно-курортное лечение', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
        [
            Paragraph('6', styleT),
            Paragraph('Диетпитание', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
        [
            Paragraph('7', styleT),
            Paragraph('Взято на диспансерное наблюдение', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
        [
            Paragraph('8', styleT),
            Paragraph('Направлено на медико-социальную экспертизу', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
            Paragraph(' ', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(10 * mm, 90 * mm, 25 * mm, 25 * mm, 25 * mm,), )
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    objs.append(Spacer(1, 5 * mm))
    objs.append(
    Paragraph('10. Рекомендации     работодателю:       санитарно-профилактические     и оздоровительные мероприятия и т.п.: Соблюдение режима труда и отдыха; сезонная вакцинация '
              '(грипп; клещевой энцефалит); зарядка на рабочем мест; закаливание; обеспечить возможность и оказать содействие работникам, нуждающимся в прохождении '
              'соответствующего обследования и лечения; обеспечить возможность и оказать содействие работникам , нуждающимся в санаторно – курортном лечении, в прохождении '
              'соответствующего СКЛ; обеспечить возможность и оказать содействие работникам, нуждающимся диспансерном наблюдении, в прохождении соответствующего наблюдения, '
              'обеспечить соблюдение санитарно- гигиенических норм условий труда.', style))



    def first_pages(canvas, document):
        canvas.saveState()
        canvas.restoreState()

    def later_pages(canvas, document):
        canvas.saveState()
        canvas.restoreState()

    doc.build(objs, onFirstPage=first_pages, onLaterPages=later_pages)

    pdf = buffer.getvalue()
    buffer.seek(0)

    today = datetime.datetime.now()
    date_now1 = datetime.datetime.strftime(today, "%y%m%d%H%M%S%f")[:-3]
    # date_now_str = str(client_id) + str(date_now1)
    # date_now_str = str(client_id) + str(date_now1)
    date_now_str = str(date_now1)
    dir_param = SettingManager.get("dir_param", default='/tmp', default_type='s')
    file_dir = os.path.join(dir_param, date_now_str + '_dir.pdf')

    def save(form, filename: str):
        with open(filename, 'wb') as f:
            f.write(form.read())

    save(buffer, filename=file_dir)

    from pdf2docx import Converter

    # pdf_file = '/Users/sergejkasanenko/Documents/tmp/results.pdf'
    docx_file = os.path.join(dir_param, date_now_str + '_dir.docx')

    # convert pdf to docx
    cv = Converter(file_dir)
    cv.convert(docx_file, start=0, end=None)
    cv.close()
    buffer.close()
    os.remove(file_dir)
    doc = docx.Document(docx_file)
    os.remove(docx_file)
    buffer.close()

    return doc


def add_needs_text(objs, text, styleT):
    objs.append(Spacer(1, 7 * mm))
    opinion = [
        [
            Paragraph(f'{text}', styleT),
            Paragraph('', styleT),
            Paragraph('', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=(125 * mm, 25 * mm, 25 * mm))
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                # ('LEFTPADDING', (1, 0), (-1, -1), 2 * mm),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)
    return objs


def add_fio_spec_diagnoz(objs, styleT):
    objs.append(Spacer(1, 0 * mm))
    opinion = [
        [
            Paragraph('№', styleT),
            Paragraph('Ф.И.О.', styleT),
            Paragraph('Подразделение', styleT),
            Paragraph('Специалист', styleT),
            Paragraph('Диагноз', styleT),
        ]
    ]
    tbl = Table(
        opinion,
        colWidths=(
            10 * mm,
            90 * mm,
            25 * mm,
            25 * mm,
            25 * mm,
        ),
    )
    tbl.setStyle(
        TableStyle(
            [
                ('GRID', (0, 0), (-1, -1), 0.75, colors.black),
                # ('LEFTPADDING', (1, 0), (-1, -1), 2 * mm),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]
        )
    )
    objs.append(tbl)

    return objs


def all_and_women(objs, styleT):
    col_width = (150 * mm, 20 * mm)
    opinion = [
        [
            Paragraph('всего,', styleT),
            Paragraph('18', styleT),
        ],
        [
            Paragraph('в том числе женщин,', styleT),
            Paragraph('6', styleT),
        ],
    ]
    tbl = Table(opinion, colWidths=col_width)
    tbl.setStyle(
        TableStyle(
            [
                ('LEFTPADDING', (1, 0), (-1, -1), 1.5 * mm),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 1.5, colors.black),
            ]
        )
    )
    objs.append(tbl)
    return objs
